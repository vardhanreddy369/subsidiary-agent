from docx import Document
from docx.shared import Pt
import datetime

doc = Document()

# Add Title
title = doc.add_heading('Initial Progress Update: Subsidiary Timeline Extraction', 0)
title.alignment = 1 # Center

# Add Date and To/From
p = doc.add_paragraph()
p.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}\n").bold = True
p.add_run("To: Dr. Christo Pirinsky, Dr. Vladimir Gatchev, Dr. Rodney Ndum\n").bold = True
p.add_run("From: Sri Vardhan Reddy Gutta, Sarayu Panditi").bold = True

doc.add_paragraph("Dear Dr. Pirinsky, Dr. Gatchev, and Dr. Ndum,")

doc.add_paragraph(
    "We wanted to provide a quick update on the subsidiary dataset you shared with us. "
    "We have successfully ported the SAS file into our data environment and have begun testing the automated "
    "extraction scripts to isolate the TimeIn, TimeOut, and source variables."
)

doc.add_heading('What We Have Accomplished So Far', level=2)
ul = doc.add_paragraph(style='List Bullet')
ul.add_run("Data Exploration: ").bold = True
ul.add_run("During our initial data processing, we found that the dataset contains exactly 436,948 unique subsidiary records.")
ul2 = doc.add_paragraph(style='List Bullet')
ul2.add_run("Automated Extraction Prototype: ").bold = True
ul2.add_run("We built a working prototype of the automated web-search script (using large language models and search APIs) "
            "as discussed. We tested it on a small sample, and it successfully parses financial news articles and SEC filings "
            "to find the exact dates for incorporation, M&A activity, or corporate restructuring.")

doc.add_heading('The Scaling Challenge', level=2)
doc.add_paragraph(
    "While the automated web-scraping script works exceptionally well on individual companies, running live searches for nearly "
    "440,000 corporate records sequentially presents a practical challenge. Executing this at scale would require significant processing "
    "time and incur high API compute costs due to the sheer volume of web requests."
)

doc.add_heading('Proposed Strategy', level=2)
doc.add_paragraph(
    "To extract this timeline data efficiently while maintaining academic rigor, we recommend a two-step approach:"
)
ol1 = doc.add_paragraph(style='List Number')
ol1.add_run("Bulk Database Matching: ").bold = True
ol1.add_run("We first cross-reference the 437,000 records against structured historical M&A databases "
            "(such as Compustat or historical SEC EDGAR filings). This standard data merge should instantly match the "
            "TimeIn dates for the vast majority of the dataset.")
ol2 = doc.add_paragraph(style='List Number')
ol2.add_run("Targeted Web Scraping for Edge Cases: ").bold = True
ol2.add_run("For the remaining ambiguous records—and specifically for tracking divestitures (TimeOut)—we will deploy our "
            "automated web-search script to parse the missing data from public sources.")

doc.add_heading('Next Steps', level=2)
doc.add_paragraph(
    "Sarayu and I are currently investigating our access to Compustat and Capital IQ via the university library to pull the bulk M&A data. "
    "Simultaneously, we are refining our matching logic to merge this bulk data safely with your dataset."
)
doc.add_paragraph(
    "Please let us know if you approve of this methodology, or if you prefer we generate a small 500-row sample using just "
    "the web-search script for your review first."
)

doc.add_paragraph("Best regards,\nSri & Sarayu")

# Save
doc.save('Update_Report.docx')
print("Saved Update_Report.docx")
